import numpy as np
import pandas as pd
import os
import tkinter  
from tkinter import Label, filedialog as tkFileDialog
from tkinter import messagebox
from scipy.stats import f
import docx
import sys

# excelファイルもしくは、CSVファイルの選択
root = tkinter.Tk()  
root.withdraw()
messagebox.showinfo("Select",".xlsx .xls .csvのどれかを選択してください") 
root.attributes("-topmost", True)
root.focus_force()
fTyp = [("",".xlsx"),("",".csv"),("",".xls")]
download_folder=os.path.join(os.environ["USERPROFILE"],"Downloads")
file_path = tkFileDialog.askopenfilename(filetypes=fTyp,initialdir=download_folder)
file_name = os.path.splitext(os.path.basename(file_path))[0]
file_type = os.path.splitext(file_path)[1][1:]
root.destroy()

# folderが無ければ作成
if not os.path.exists(download_folder+"/result_ttest"):
    os.makedirs(download_folder+"/result_ttest")
if not os.path.exists(download_folder+"/png"):
    os.makedirs(download_folder+"/png")



################################
# 関数
################################

# 引数が数字かどうか
def is_num(x):
    try:
        float(x)
    except ValueError:
        return False
    else:
        return True

# t検定の対応ありなしの選択
def is_pair():
    global bol1
    bol1 = False
    root = tkinter.Tk()
    a=int(root.winfo_screenwidth() /2 -150)
    b=int(root.winfo_screenheight() /2 -100)
    root.geometry('{}x{}+{}+{}'.format(300, 200, a, b))
    root.title('ｔ検定')
    root.attributes("-topmost", True)
    # ラジオボタンのラベルをリスト化する
    rdo_txt = ['対応あり','対応なし']
    # ラジオボタンの状態
    rdo_var = tkinter.IntVar()
    for i in range(len(rdo_txt)):
        rdo = tkinter.Radiobutton(root, value=i, variable=rdo_var, text=rdo_txt[i]) 
        rdo.place(x=50, y=30 + (i * 24))
    # ボタンクリックイベント
    def btn_click():
        global bol1
        num =rdo_var.get()
        if(num==0):
            bol1=True
        else:
            bol1 = False
        global root
        root.quit()
    # ボタン作成 
    btn = tkinter.Button(root, text='決定', command=btn_click)
    btn.place(x=100, y=150)    
    root.mainloop()
    root.destroy()
    return bol1

# 両側検定か片側検定かの選択
def var_select(txt2):
    global alternative
    alternative = ""
    root = tkinter.Tk()
    a=int(root.winfo_screenwidth() /2 -200)
    b=int(root.winfo_screenheight() /2 -100)
    root.geometry('{}x{}+{}+{}'.format(400, 200, a, b))
    root.title('alternativeの設定')
    root.attributes("-topmost", True)
    # ラジオボタンのラベルをリスト化する
    rdo_txt=[
        "両側検定",
        "片側検定less 対立仮説：(列"+txt2[0]+" ＜ 列"+txt2[1]+")",
        "片側検定greatr 対立仮説： (列"+txt2[0]+" ＜ 列"+txt2[1]+")",
        "全部出力"
    ]

    
    # ラジオボタンの状態
    rdo_var = tkinter.IntVar()
    for i in range(len(rdo_txt)):
        rdo = tkinter.Radiobutton(root, value=i, variable=rdo_var, text=rdo_txt[i]) 
        rdo.place(x=50, y=30 + (i * 24))
    # ボタンクリックイベント
    def btn_click2():
        global alternative
        num = rdo_var.get()
        if(num==0):
            alternative="two-sided"
        elif(num==1):
            alternative="less"
        elif(num==2):
            alternative="greater"
        elif(num==3):
            alternative='all'
        global root
        root.quit()
    # ボタン作成 
    btn = tkinter.Button(root, text='決定', command=btn_click2)
    btn.place(x=100, y=150)    
    root.mainloop()
    root.destroy()
    return alternative

def ftest(a, b):
    #　統計量Fの計算
    v1 = np.var(a, ddof=1)
    v2 = np.var(b, ddof=1)
    n1 = len(a)
    n2 = len(b)
    f_value = v1/v2
    df1 = n1 -1
    df2 = n2 -1

    # 帰無仮説が正しい場合にFが従う確率分を生成
    f_frozen = f.freeze(dfn=n1-1, dfd=n2-1)

    # 右側
    p1 = f_frozen.sf(f_value)
    # 左側
    p2 = f_frozen.cdf(f_value)
    # 小さい方の2倍がp値
    p_value = min(p1, p2) * 2

    # 統計量Fとp値を返す
    return f_value, p_value,df1,df2

##############################################
##############################################


# ファイルがExcelのとき
if(file_type=="xlsx" or file_type=="xls"):
    try:
        df_main = pd.read_excel(file_path,sheet_name=0,index_col=None,header=None)
    except UnicodeDecodeError:
        messagebox.showerror("Error","ファイルが文字化けしています。") 
        sys.exit()

# ファイルがCSVのとき
elif(file_type=="csv"):
    try:
        df_main = pd.read_csv(file_path,index_col=None,header=None, encoding="cp932")
    except UnicodeDecodeError:
        messagebox.showerror("Error","ファイルが文字化けしています。") 
        sys.exit()


# カラム名が設定されているとき
if(is_num(df_main.iat[0,0])==False):
    df_main.columns = df_main.iloc[0].tolist()
    df_main = df_main.drop(df_main.index[0])
# 指定されていないとき
else:
    df_main.columns = ["Column1","Column2"]

# データが2列のデータか確認
if(len(df_main.columns)<2 or len(df_main.columns) >2):
    messagebox.showerror("Error","ファイルのデータが、2列のデータを選択してください")
    sys.exit()

# データに数値以外があるかどうかを確認
try:
    df_main = df_main.astype("float")
except ValueError:
    messagebox.showerror("TypeError","数値(int,float)以外のデータが含まれています")
    sys.exit()

# 2標本のデータラベル 
list_col = df_main.columns.tolist()
# 値をリストで取得
list1 = df_main.iloc[:,0].tolist()
list2 = df_main.iloc[:,1].tolist()

# t検定の詳細設定
is_paired_bool=is_pair()
alternative_str = var_select(df_main.columns.tolist())

import scipy.stats as stats
import seaborn as sns
import matplotlib.pyplot as plt
import statistics
import japanize_matplotlib

# Wordを新規作成
doc=docx.Document()
doc.add_heading("2標本の検定手順について",0)
doc.add_paragraph("検定は以下の手順で行っています。")
doc.add_paragraph("     １．データが正規分布に従うかの検定（シャピロ・ウィルク検定とQ-Qプロットによる確認）")
doc.add_paragraph("     ２．2標本の母分散が等しいかどうかの検定（F検定）")
doc.add_paragraph("     ３．2標本の母平均が等しいかどうかの検定")
doc.add_paragraph("\n\n元データ")
doc.add_paragraph(str(df_main))
doc.add_page_break()


# 正規性の検定（Q-Qプロット）
doc.add_heading("１．データが正規分布に従うかの検定",0)
doc.add_heading("データの正規性をQ-Qプロットによる可視化で確認",1)
img1_path = download_folder+'/png/qq1_py.png'
img2_path = download_folder+'/png/qq2_py.png'
plt.figure(figsize=(4.8,3.6))
stats.probplot(list1, dist="norm", plot=plt)
plt.title(str(list_col[0]) + "  Q-Q Plot",fontname="MS Gothic")
plt.savefig(img1_path)
plt.close()

plt.figure(figsize=(4.8,3.6))
stats.probplot(list2, dist="norm", plot=plt)
plt.title(str(list_col[1]) + "  Q-Q Plot",fontname="MS Gothic")
plt.savefig(img2_path)
plt.close()
doc.add_picture(img1_path)
doc.add_picture(img2_path)
doc.add_page_break()

# シャピロ・ウィルク検定による正規性の検定
doc.add_heading("シャピロ・ウィルク検定による正規性の検定",1)
doc.add_paragraph("帰無仮説：データの母集団が正規分布に従う")

doc.add_paragraph("・"+list_col[0] + "のシャピロ・ウィルク検定の結果")
W_1, p_1 = stats.shapiro(list1)
W_1_str=str(round(W_1,4))
p_1_str=str(round(p_1,4))
doc.add_paragraph("--------------------------------------")
doc.add_paragraph("W=\t"+W_1_str)
doc.add_paragraph("P値=\t"+p_1_str)
doc.add_paragraph("--------------------------------------")
if(p_1 < 0.05):
    doc.add_paragraph("帰無仮説が有意水準が5%で棄却され,この検定では"+ list_col[0] +"が正規分布に従うとは言えない")   
    sha_result_bool1 =False
    sha_result_text1 = "正規性はない"
else:
    doc.add_paragraph(list_col[0]+"は正規分布に従うといえる")
    sha_result_bool1 =True
    sha_result_text1 = "正規性がある"

doc.add_paragraph("\n\n・"+list_col[1] + "のシャピロ・ウィルク検定の結果")
W_2, p_2 = stats.shapiro(list2)
W_2_str=str(round(W_2,4))
p_2_str=str(round(p_2,4))
doc.add_paragraph("--------------------------------------")
doc.add_paragraph("W=\t"+W_2_str)
doc.add_paragraph("P値=\t"+p_2_str)
doc.add_paragraph("--------------------------------------")
if(p_2 < 0.05):
    doc.add_paragraph("帰無仮説が有意水準が5%で棄却され,この検定では"+ list_col[1] +"が正規分布に従うとは言えない")   
    sha_result_bool2 =False
    sha_result_text2 = "正規性はない"

else:
    doc.add_paragraph(list_col[1]+"は正規分布に従うといえる")
    sha_result_bool2 =True
    sha_result_text2 = "正規性がある"
doc.add_page_break()

#等分散性の検定
doc.add_heading("２．2標本の母分散が等しいかどうかの検定",0)
doc.add_heading("F検定",1)
doc.add_paragraph("帰無仮説：２群間の母分散は等しい")
doc.add_paragraph("F検定の結果")
doc.add_paragraph("--------------------------------------")
f_value,p_value,df1,df2=ftest(list1,list2)
doc.add_paragraph("F統計量=\t"+str(round(f_value,4)))
doc.add_paragraph("P値=\t"+str(round(p_value,4)))
doc.add_paragraph(list_col[0]+"の自由度=\t"+str(df1))
doc.add_paragraph(list_col[1]+"の自由度=\t"+str(df2))
doc.add_paragraph("--------------------------------------")
if(p_value<0.05):
    doc.add_paragraph("帰無仮説が有意水準が5%で棄却され,この検定では等分散であるとは言えない")
    EV_bool=False
    EV_text = "（F検定の結果、P値が"+str(round(p_value,4))+"より、等分散でないとする）"
else:
    doc.add_paragraph("2標本は等分散であるといえる")
    EV_bool=True
    EV_text="（F検定の結果、P値が"+str(round(p_value,4))+"より、等分散であるといえる）"
doc.add_page_break()

doc.add_heading("データの等分散性をバイオリンプロットとヒストグラムによる確認",1)
doc.add_paragraph("・バイオリンプロット")
sns.violinplot(x=[list_col[0] for i in range(len(list1))]+[list_col[1] for i in range(len(list2))],
y=np.concatenate([list1, list2]))
v_img_path = download_folder+'/png/violin.png'
plt.savefig(v_img_path)
doc.add_picture(v_img_path)

# 分散を比べるヒストグラム
doc.add_page_break()
doc.add_paragraph("・ヒストグラム")
hist_img_path = download_folder+'/png/hist.png'
fig = plt.figure()
ax = fig.add_subplot(111)
ax.hist(list1, bins=15,  color='red', alpha = 0.7) 
ax.hist(list2, bins=15,  color='blue', alpha = 0.7) 
ax.set_title(list_col[0]+":Red     "+list_col[1]+":Blue") 
plt.savefig(hist_img_path)
doc.add_picture(hist_img_path)
doc.add_page_break()



from PIL import ImageTk, Image
# 正規性の確認を通知
RE_bool=False
root = tkinter.Tk()
a=int(root.winfo_screenwidth() /2 -600)
b=int(root.winfo_screenheight() /2 -350)
root.geometry('{}x{}+{}+{}'.format(1200, 700, a, b))
root.title('検定方法の最終決定')
root.attributes("-topmost", True)

sha_result1 = list_col[0]+" のシャピロウィルクの検定では、P値が"+p_1_str+\
    "で、"+sha_result_text1+"という結果です"
Label1 = tkinter.Label(root,text=sha_result1)
Label1.pack(side="top",padx=5, pady=5, anchor=tkinter.W)

sha_result2 = list_col[1]+" のシャピロウィルクの検定では、P値が"+p_2_str+\
    "で、"+sha_result_text2+"という結果です"
Label2 = tkinter.Label(root,text=sha_result2)
Label2.pack(side="top",padx=5, pady=5, anchor=tkinter.W)

Label3 = tkinter.Label(root,text=EV_text)
Label3.pack(side="top",padx=5, pady=5, anchor=tkinter.W)

# font4=font.Font(size=20)
text4="下のQQプロットの結果とシャピロ・ウィルクの検定結果から、データの正規性があるとみなしますか？"
Label4 = tkinter.Label(root,text=text4)
Label4.pack(side="top",padx=5, pady=5, anchor=tkinter.W)


# ラジオボタンのラベルをリスト化する
if(is_paired_bool ==True):   

    rdo_txt=[
        "正規性があるとみなす。(対応のあるｔ検定)",
        "正規性がないとみなす。(Wilcoxonの符号付き順位和検定)"
    ]
else:

    if(EV_bool==True):
        rdo_txt=[
            "正規性があるとみなす。(スチューデントのｔ検定)",
            "正規性がないとみなす。(Wilcoxonの順位和検定)"
        ]
    else:
        rdo_txt=[
            "正規性があるとみなす。(Welchのｔ検定)",
            "正規性がないとみなす。(Wilcoxonの順位和検定)"
        ]
# ラジオボタンの状態
rdo_var1 = tkinter.IntVar()
rdo = tkinter.Radiobutton(root, value=0, variable=rdo_var1, text=rdo_txt[0]) 
rdo.pack(side="top")
rdo2 = tkinter.Radiobutton(root, value=1, variable=rdo_var1, text=rdo_txt[1]) 
rdo2.pack(side="top")

# ボタンクリックイベント
def btn_click4():
    global RE_bool
    num = labelValue.cget("text")
    if(num==0):
        RE_bool=True
    else:
        RE_bool=False
    global root
    root.quit()
# ボタン作成 
btn = tkinter.Button(root, text='決定', command=btn_click4)
btn.pack(side="top")  

labelValue = tkinter.Label(root, textvariable=rdo_var1)
# labelValue.pack(side="top")

canvas1 = tkinter.Canvas(root,bg="black", width=480, height=360)
canvas1.pack(fill = 'x', padx=20, side = 'left')
image1=Image.open(img1_path)
photo1 = ImageTk.PhotoImage(image1,master=root)
canvas1.create_image(0,0,anchor='nw',image=photo1)

canvas2 = tkinter.Canvas(root,bg="black", width=480, height=360)
canvas2.pack(fill = 'x', padx=20, side = 'left')
image2=Image.open(img2_path)
photo2 = ImageTk.PhotoImage(image2,master=root)
canvas2.create_image(0,0,anchor='nw',image=photo2)


root.mainloop()
root.destroy()


################
# t検定
################
result_list=[]
Ah = ["two-sided","less","greater"]
# 対応のあるT検定
if(is_paired_bool==True):
    # 正規性がある場合(対応のあるT検定)
    if(RE_bool==True ):
        doc.add_heading("３．対応あり、正規性あり\n「対応のあるT検定」",0)
        doc.add_paragraph("Rでのコマンド(両側検定):t.test(list1,list2,paired=T,alternative='two.sided')")
        if(alternative_str=='all'):
            for i in Ah:
                statistic,Pvalue = stats.ttest_rel(list1,list2,alternative=i)
                result_list.append([statistic,Pvalue])
        else:
            statistic,Pvalue = stats.ttest_rel(list1,list2,alternative=alternative_str)
    # 正規性がない場合(Wilcoxonの符号付き順位和検定)
    else:
        doc.add_heading("３．対応あり、正規性なし\n「Wilcoxonの符号付き順位和検定(連続性の補正)」",0)
        doc.add_paragraph("Rでのコマンド(両側検定):wilcox.test(list1,list2,paired=T,exact=F,alternative='two.sided')")
        if(alternative_str=='all'):
            for i in Ah:
                statistic,Pvalue = stats.wilcoxon(list1,list2,alternative=i,correction=True,mode="approx")
                result_list.append([statistic,Pvalue])
        else:
            # modeを"approx"にするとR言語と同じ結果が得られるが、正確な"exact"を採用した
            statistic,Pvalue= stats.wilcoxon(list1,list2,alternative=alternative_str,correction=True,mode="approx")
# 対応のないT検定
else:
    # 正規性がある場合
    if(RE_bool==True ):
        # 等分散である場合（スチューデントのｔ検定）
        if(EV_bool==True):
            doc.add_heading("３．対応なし、正規性あり、等分散性あり\n「スチューデントのT検定」",0)
            doc.add_paragraph("Rでのコマンド(両側検定):t.test(list1,list2,var.equal=T,alternative='two.sided')")
            if(alternative_str=='all'):
                for i in Ah:
                    statistic,Pvalue= stats.ttest_ind(list1, list2, equal_var=True,alternative=i)
                    result_list.append([statistic,Pvalue])
            else:
                statistic,Pvalue= stats.ttest_ind(list1, list2, equal_var=True,alternative=alternative_str)
        # 等分散でない場合(WelchのT検定)
        else:
            doc.add_heading("３．対応なし、正規性あり、等分散性なし\n「WelchのT検定」",0)
            doc.add_paragraph("Rでのコマンド(両側検定):t.test(list1,list2,var.equal=F,alternative='two.sided')")
            if(alternative_str=='all'):
                for i in Ah:
                    statistic,Pvalue=stats.ttest_ind(list1, list2, equal_var=False,alternative=i)
                    result_list.append([statistic,Pvalue])
            else:
                statistic,Pvalue=stats.ttest_ind(list1, list2, equal_var=False,alternative=alternative_str)
    # 正規性がない場合（Wilcoxonの順位和検定）
    else:
        doc.add_heading("３．対応なし、正規性なし\n「Wilcoxonの順位和検定」(連続性の補正)",0)
        doc.add_paragraph("Rでのコマンド(両側検定):wilcox.test(list1,list2,exact=F,alternative='two.sided')")
        if(alternative_str=='all'):
            for i in Ah:
                    statistic,Pvalue=stats.mannwhitneyu(list1,list2, alternative=i)
                    result_list.append([statistic,Pvalue])
        else:
            statistic,Pvalue=stats.mannwhitneyu(list1,list2, alternative=alternative_str)

def result_p(alternative,pval):
    global doc
    if(alternative=="two-sided"):
        doc.add_paragraph("帰無仮説：2郡間の平均値に差がない")
        if(pval<0.05):
            doc.add_paragraph("帰無仮説が有意水準が5%で棄却され、2郡間の平均値に差があるといえる")
        elif(pval<0.10):
            doc.add_paragraph("2郡間の平均値に有意傾向があるといえる")
        else:
            doc.add_paragraph("2郡間の平均値に有意な差はない")
    elif(alternative=="less"):
        doc.add_paragraph("帰無仮説：列"+list_col[0]+"の平均値が、列"+list_col[1]+"の平均値よりも大きい")
        if(pval<0.05):
            doc.add_paragraph("帰無仮説が有意水準が5%で棄却され、対立仮説が成り立つといえる")
        elif(pval<0.10):
            doc.add_paragraph("有意傾向があるといえる")
        else:
            doc.add_paragraph("帰無仮説は棄却されない")     
    elif(alternative=="greater"):
        doc.add_paragraph("帰無仮説：列"+list_col[0]+"の平均値が、列"+list_col[1]+"の平均値よりも小さい")
        if(pval<0.05):
            doc.add_paragraph("帰無仮説が有意水準が5%で棄却され、対立仮説が成り立つといえる")
        elif(pval<0.10):
            doc.add_paragraph("有意傾向があるといえる")
        else:
            doc.add_paragraph("帰無仮説は棄却されない")     
      




Bh = [
    "両側検定",
    "片側検定less 対立仮説：(列"+list_col[0]+" ＜ 列"+list_col[1]+")",
    "片側検定greater 対立仮説：(列"+list_col[0]+" ＞ 列"+list_col[1]+")"
]
doc.add_paragraph("--------------------------------------")
if(alternative_str=='all'):
    for i in range(0,3):
        doc.add_paragraph(Bh[i]+"の結果")
        if(RE_bool==False):
            doc.add_paragraph("W=\t"+str(round(result_list[i][0],4)))
        else:
            doc.add_paragraph("T値=\t"+str(round(result_list[i][0],4)))
        doc.add_paragraph("P値=\t"+str(round(result_list[i][1],4)))
        result_p(Ah[i],result_list[i][1])
        doc.add_paragraph("--------------------------------------")
else:
    doc.add_paragraph(Bh[Ah.index(alternative_str)]+"の結果")
    if(RE_bool==False):
        doc.add_paragraph("W=\t"+str(round(statistic,4)))
    else:
        doc.add_paragraph("T値=\t"+str(round(statistic,4)))
    doc.add_paragraph("P値=\t"+str(round(Pvalue,4)))
    result_p(alternative_str,Pvalue)
    doc.add_paragraph("--------------------------------------")


mean_list=[statistics.mean(list1),statistics.mean(list2)]
var_list=[np.var(list1,ddof=1),np.var(list2,ddof=1)]
std_list=[np.std(list1,ddof=1),np.std(list2,ddof=1)]

doc.add_page_break()
doc.add_heading("グラフ化",1)
doc.add_paragraph(list_col[0]+"の平均値＝"+str(mean_list[0]))
doc.add_paragraph(list_col[1]+"の平均値＝"+str(mean_list[1]))
doc.add_paragraph(list_col[0]+"の不偏分散＝"+str(var_list[0]))
doc.add_paragraph(list_col[1]+"の不偏分散＝"+str(var_list[1]))
doc.add_paragraph(list_col[0]+"の不偏標準偏差＝"+str(std_list[0]))
doc.add_paragraph(list_col[1]+"の不偏標準偏差＝"+str(std_list[1]))

# 棒グラフ化
label_x=list_col
n=len(list1)
x=[1,2]
max_mean=((max(mean_list[0],mean_list[1])+max(std_list[0],std_list[1]))*1.3)
fig1,ax1=plt.subplots(figsize=(3,4))
ax1.bar(x,mean_list,width=0.5,color="#87ceeb",tick_label=label_x,align="center",yerr=std_list,ecolor="black",capsize=3)
step=0
if(max_mean<10):
    step=0.5
elif(max_mean<20):
    step=1
elif(max_mean<40):
    step=2
elif(max_mean<60):
    step=3
elif(max_mean<100):
    step=5
else:
    step=int((max_mean/20)/10)*10

if(step!=0):
    ax1.set_yticks(np.arange(0,max_mean+1,step=step))
else:
    ax1.set_yticks(np.arange(0,max_mean+1,step=(max_mean/20)))

ax1.set_xlabel("+p < .10, *p < .05, **p < .01, ***p < .001")
ax1.text(0.7,max_mean/1.15,"mean+SD\nN="+str(n),fontsize=7,verticalalignment="baseline",horizontalalignment="left")
if(max(len(label_x[0]),len(label_x[1]))>8):
    plt.xticks(rotation=45)
plt.tight_layout()
# 有意差のあるものに＊を加える
arrow_dict = dict(arrowstyle="-",color="black", connectionstyle ="bar,fraction=0.1")
if(len(result_list)>0):
    p_value=result_list[0][1]
    doc.add_paragraph("alternativeを3つに設定した場合、グラフの有意差は、両側検定のものをもとに作成しています")
else:
    p_value=Pvalue
text=""
if(p_value<0.001):
    text="***"
elif(p_value<0.01):
    text="**"
elif(p_value<0.05):
    text="*"
elif(p_value<0.1):
    text="+"
else:
    text="NS"
    # 有意であれば棒グラフに＊を追加する
if(text != "NS"):
    # 追加位置を求める
    pos_y=(max_mean/1.3)+(step/5)
    pos_y2=(max_mean/1.3)
    # [をついか
    ax1.annotate("",xy = (1,pos_y),xytext=(2,pos_y),color="black", arrowprops = arrow_dict)
    ax1.text(1.5,pos_y2+step,text,size=20,horizontalalignment="center")
img_ttest_path=download_folder+"/png/ttest_bar.png"
plt.savefig(img_ttest_path)
plt.close()
doc.add_picture(img_ttest_path)



# 保存のときのファイル名を取得
import tkinter.ttk as ttk
save_file_name =""
#ボタンがクリックされたときの関数
def clicked():
    global save_file_name
    save_file_name = entry.get()
    root.quit()

root = tkinter.Tk()  
root.lift()
a=int(root.winfo_screenwidth() /2 -150)
b=int(root.winfo_screenheight() /2 -100)
root.geometry('{}x{}+{}+{}'.format(300, 200, a, b))
root.title("保存ファイル名の入力")
root.attributes("-topmost", True)
# フレームの作成と設置
frame = ttk.Frame(root)
frame.grid(column=0, row=0, sticky=tkinter.NSEW, padx=5, pady=10)

label = ttk.Label(frame, text="ファイル名：")
entry = ttk.Entry(frame)
button_execute = ttk.Button(frame, text="保存", command=clicked)

# 各種ウィジェットの設置
label.grid(row=0, column=0)
entry.grid(row=0, column=1)
button_execute.grid(row=1, column=1)
# Entryウィジェットへ文字列のセット
entry.insert(tkinter.END,file_name + "_ttest")
root.mainloop()
root.destroy()

# pigフォルダの削除
import shutil
shutil.rmtree(download_folder + "/png/")

mess = download_folder+"/result_ttest/" +save_file_name+".docx"
try:
    doc.save(mess)
except PermissionError:
    messagebox.showerror("Error","同じ名前のファイルが開かれています。\nファイルを閉じてからもう一度実行しなおしてください")
    sys.exit()
messagebox.showinfo("処理完了",mess+"に結果が保存されました")